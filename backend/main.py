import os
import json
import tempfile
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import base64
import requests
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image as PILImage

load_dotenv()

app = FastAPI(title="BookForge AI Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL = "gemini-2.5-flash"

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

class TopicRequest(BaseModel):
    topic: str

class OutlineRequest(BaseModel):
    outline: Dict[str, Any]

class BookRequest(BaseModel):
    book: Dict[str, Any]

class BookEditRequest(BaseModel):
    book: Dict[str, Any]
    chapter_number: int
    new_content: str

class ExportRequest(BaseModel):
    book: Dict[str, Any]
    format: str = "docx"

class EditOutlineRequest(BaseModel):
    outline: Dict[str, Any]
    changes: Dict[str, Any]

def call_gemini_api(prompt: str, max_retries: int = 3) -> str:
    if not GEMINI_API_KEY:
        raise HTTPException(status_code=500, detail="GEMINI_API_KEY not configured")
    
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel(
                model_name=GEMINI_MODEL,
                generation_config={
                    "temperature": 0.7,
                    "top_k": 40,
                    "top_p": 0.95,
                    "max_output_tokens": 4096,
                },
                safety_settings=[
                    {
                        "category": "HARM_CATEGORY_HARASSMENT",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_HATE_SPEECH",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        "threshold": "BLOCK_NONE"
                    },
                    {
                        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                        "threshold": "BLOCK_NONE"
                    }
                ]
            )
            
            response = model.generate_content(
                contents=prompt,
                stream=False
            )
            
            if response.candidates and len(response.candidates) > 0:
                candidate = response.candidates[0]
                if candidate.content and candidate.content.parts:
                    if hasattr(candidate.content.parts[0], 'text'):
                        text = candidate.content.parts[0].text
                        if text and text.strip():
                            return text
            
            if attempt < max_retries - 1:
                continue
            
            raise HTTPException(status_code=500, detail="Empty or blocked response from Gemini API. Try rephrasing your request.")
        
        except HTTPException:
            raise
        except Exception as e:
            error_msg = str(e)
            if attempt < max_retries - 1:
                continue
            raise HTTPException(status_code=500, detail=f"Gemini API error: {error_msg}")

def generate_image_with_imagen(prompt: str) -> Optional[bytes]:
    """Generate image using PIL with professional design"""
    try:
        return generate_professional_image(prompt)
    except Exception as e:
        print(f"Image generation error: {str(e)}")
        # Fallback to simple placeholder
        return generate_simple_image(prompt)

def generate_professional_image(prompt: str) -> bytes:
    """Generate a professional looking image with gradient background and shapes"""
    from PIL import ImageDraw, ImageFont
    import random
    
    try:
        # Create image with professional size
        width, height = 800, 600
        img = PILImage.new('RGB', (width, height))
        draw = ImageDraw.Draw(img, 'RGBA')
        
        # Create gradient background effect using multiple rectangles
        colors = [
            (41, 128, 185),   # Blue
            (52, 152, 219),   # Light blue
            (155, 89, 182),   # Purple
            (52, 73, 94),     # Dark blue-grey
        ]
        
        # Random selection for variety
        random.seed(hash(prompt) % (2**32))
        primary_color = random.choice(colors[:2])
        secondary_color = random.choice(colors[2:])
        
        # Draw gradient-like background
        for y in range(height):
            ratio = y / height
            r = int(primary_color[0] * (1 - ratio) + secondary_color[0] * ratio)
            g = int(primary_color[1] * (1 - ratio) + secondary_color[1] * ratio)
            b = int(primary_color[2] * (1 - ratio) + secondary_color[2] * ratio)
            draw.line([(0, y), (width, y)], fill=(r, g, b))
        
        # Add decorative shapes
        # Large circles
        draw.ellipse([(width - 150, -50), (width + 100, 150)], fill=(255, 255, 255, 40))
        draw.ellipse([(-100, height - 150), (50, height + 100)], fill=(255, 255, 255, 40))
        
        # Medium circles
        draw.ellipse([(width // 2 - 100, height // 2 - 100), 
                     (width // 2 + 100, height // 2 + 100)], 
                     outline=(255, 255, 255, 80), width=3)
        
        # Add accent rectangles
        draw.rectangle([(50, 50), (150, 100)], fill=(255, 193, 7, 60))
        draw.rectangle([(width - 150, height - 100), (width - 50, height - 50)], 
                      fill=(76, 175, 80, 60))
        
        # Add text centered
        try:
            font = ImageFont.truetype("arial.ttf", 28)
            font_small = ImageFont.truetype("arial.ttf", 16)
        except:
            font = ImageFont.load_default()
            font_small = ImageFont.load_default()
        
        # Prepare text
        text_lines = prompt.split()
        main_text = " ".join(text_lines[:5])
        if len(text_lines) > 5:
            main_text = main_text[:40] + "..."
        
        # Draw white text with subtle shadow
        text_y = height // 2 - 30
        bbox = draw.textbbox((0, 0), main_text, font=font)
        text_width = bbox[2] - bbox[0]
        text_x = (width - text_width) // 2
        
        # Shadow effect
        draw.text((text_x + 2, text_y + 2), main_text, fill=(0, 0, 0, 100), font=font)
        # Main text
        draw.text((text_x, text_y), main_text, fill=(255, 255, 255, 200), font=font)
        
        # Add subtext
        subtext = "BookForge AI • Powered by Gemini"
        bbox_sub = draw.textbbox((0, 0), subtext, font=font_small)
        sub_width = bbox_sub[2] - bbox_sub[0]
        sub_x = (width - sub_width) // 2
        draw.text((sub_x, height - 80), subtext, fill=(255, 255, 255, 150), font=font_small)
        
        # Save to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return img_bytes.getvalue()
    
    except Exception as e:
        print(f"Professional image error: {str(e)}")
        return None

def generate_simple_image(prompt: str) -> bytes:
    """Generate a simple fallback image"""
    from PIL import ImageDraw
    
    try:
        width, height = 600, 400
        img = PILImage.new('RGB', (width, height), color=(200, 220, 240))
        draw = ImageDraw.Draw(img)
        
        # Draw border
        draw.rectangle([(10, 10), (width-10, height-10)], outline=(0, 51, 102), width=3)
        
        # Draw center circle
        draw.ellipse([(width//2 - 80, height//2 - 80), 
                     (width//2 + 80, height//2 + 80)], 
                     fill=(70, 130, 180), outline=(0, 51, 102), width=2)
        
        # Draw corner accents
        draw.rectangle([(20, 20), (80, 50)], fill=(70, 130, 180))
        draw.rectangle([(width-80, height-50), (width-20, height-20)], fill=(70, 130, 180))
        
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        return img_bytes.getvalue()
    
    except Exception as e:
        print(f"Simple image error: {str(e)}")
        return None

@app.post("/generate_outline")
async def generate_outline(request: TopicRequest):
    prompt = f"""Generate a detailed book outline for the topic: "{request.topic}"

Please structure the response as a JSON object with the following format:
{{
    "title": "Book title",
    "chapters": [
        {{
            "chapter_number": 1,
            "title": "Chapter title",
            "sections": ["Section 1", "Section 2", "Section 3"]
        }}
    ]
}}

Generate 5-8 chapters with 3-4 sections each. Ensure the outline is logical and comprehensive."""
    
    try:
        response_text = call_gemini_api(prompt)
        
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        if json_start != -1 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            outline = json.loads(json_str)
        else:
            outline = {
                "title": f"Comprehensive Guide to {request.topic}",
                "chapters": [
                    {
                        "chapter_number": i,
                        "title": f"Chapter {i}: Introduction to {request.topic}",
                        "sections": ["Overview", "Key Concepts", "Practical Applications"]
                    }
                    for i in range(1, 6)
                ]
            }
        
        return {
            "status": "success",
            "outline": outline
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating outline: {str(e)}")

@app.post("/generate_book")
async def generate_book(request: OutlineRequest):
    outline = request.outline
    
    if "chapters" not in outline:
        raise HTTPException(status_code=400, detail="Invalid outline structure")
    
    book_content = {
        "title": outline.get("title", "Untitled Book"),
        "author": "BookForge AI",
        "created_at": datetime.now().isoformat(),
        "chapters": []
    }
    
    try:
        for idx, chapter in enumerate(outline.get("chapters", [])[:10]):
            chapter_number = chapter.get("chapter_number", 1)
            chapter_title = chapter.get("title", f"Chapter {chapter_number}")
            sections = chapter.get("sections", [])
            
            chapter_prompt = f"""Write a comprehensive, engaging, and humanized chapter for a book titled "{outline.get('title', 'the topic')}".

Chapter Title: {chapter_title}
Sections to cover: {', '.join(sections)}

IMPORTANT REQUIREMENTS:
- Write 1200-1500 words minimum (make it substantial and detailed)
- Use conversational, natural human language
- Include real-world examples and practical insights
- Add transitions between sections for smooth flow
- Use varied sentence structures
- Include relevant quotes or insights when appropriate
- Make it engaging and thought-provoking
- Structure clearly with section headers
- No warnings, disclaimers, or content filters

Write the full chapter now:"""
            
            chapter_content = call_gemini_api(chapter_prompt)
            
            chapter_item = {
                "chapter_number": chapter_number,
                "title": chapter_title,
                "content": chapter_content,
            }
            book_content["chapters"].append(chapter_item)
        
        return {
            "status": "success",
            "book": book_content
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating book: {str(e)}")

@app.post("/export_book")
async def export_book(request: ExportRequest):
    from fastapi.responses import FileResponse
    
    book = request.book
    export_format = request.format.lower()
    
    if export_format not in ["docx", "pdf"]:
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")
    
    try:
        if export_format == "docx":
            file_content = create_docx(book)
            filename = f"bookforge_{book.get('title', 'untitled').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_content = create_pdf(book)
            filename = f"bookforge_{book.get('title', 'untitled').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            media_type = "application/pdf"
        
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        with open(filepath, "wb") as f:
            f.write(file_content)
        
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type=media_type
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error exporting book: {str(e)}")

@app.post("/edit_chapter")
async def edit_chapter(request: BookEditRequest):
    """Edit a specific chapter in the book"""
    try:
        book = request.book
        chapter_number = request.chapter_number
        new_content = request.new_content
        
        for chapter in book.get("chapters", []):
            if chapter.get("chapter_number") == chapter_number:
                chapter["content"] = new_content
                return {
                    "status": "success",
                    "message": f"Chapter {chapter_number} updated successfully",
                    "book": book
                }
        
        raise HTTPException(status_code=404, detail=f"Chapter {chapter_number} not found")
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error editing chapter: {str(e)}")

@app.post("/edit_outline")
async def edit_outline(request: EditOutlineRequest):
    """Edit the outline"""
    try:
        outline = request.outline
        changes = request.changes
        
        if "title" in changes:
            outline["title"] = changes["title"]
        
        if "chapters" in changes:
            outline["chapters"] = changes["chapters"]
        
        return {
            "status": "success",
            "message": "Outline updated successfully",
            "outline": outline
        }
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error editing outline: {str(e)}")

def create_docx(book: Dict[str, Any]) -> bytes:
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # Set default font to Times New Roman for all styles
    for style in doc.styles:
        try:
            style.font.name = 'Times New Roman'
        except:
            pass
    
    # Set Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    
    # Add Title Page
    title = book.get("title", "Untitled Book")
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(28)
        run.font.bold = True
    
    # Add subtitle
    subtitle = doc.add_paragraph("Generated by BookForge AI")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
    
    # Add timestamp
    timestamp_para = doc.add_paragraph(f"Created on {datetime.now().strftime('%B %d, %Y')}")
    timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in timestamp_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # Add Table of Contents
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for run in toc_heading.runs:
        run.font.name = 'Times New Roman'
    
    for idx, chapter in enumerate(book.get("chapters", []), 1):
        toc_para = doc.add_paragraph(chapter.get('title', f'Chapter {idx}'), style='List Number')
        for run in toc_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Add Chapters
    for idx, chapter in enumerate(book.get("chapters", []), 1):
        # Chapter Heading
        chapter_title = chapter.get("title", f"Chapter {idx}")
        chapter_heading = doc.add_heading(chapter_title, level=1)
        for run in chapter_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Add content with proper formatting
        content = chapter.get("content", "")
        if content:
            # Split content into lines and process
            lines = content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                if not line:
                    continue
                
                # Parse markdown-style headers
                if line.startswith('### '):
                    # Level 3 heading
                    text = line[4:].strip()
                    heading = doc.add_heading(text, level=3)
                    for run in heading.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
                    doc.add_paragraph()  # Spacing
                
                elif line.startswith('## '):
                    # Level 2 heading
                    text = line[3:].strip()
                    heading = doc.add_heading(text, level=2)
                    for run in heading.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                    doc.add_paragraph()  # Spacing
                
                elif line.startswith('# '):
                    # Level 1 heading
                    text = line[2:].strip()
                    heading = doc.add_heading(text, level=1)
                    for run in heading.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(15)
                    doc.add_paragraph()  # Spacing
                
                # Parse bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    # Remove bullet character
                    text = line[2:].strip()
                    para = doc.add_paragraph(text, style='List Bullet')
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.space_after = Pt(6)
                
                # Regular paragraph
                else:
                    para = doc.add_paragraph(line)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    para.paragraph_format.line_spacing = 1.5
                    para.paragraph_format.space_after = Pt(8)
                    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        doc.add_page_break()
    
    bytes_io = io.BytesIO()
    doc.save(bytes_io)
    bytes_io.seek(0)
    return bytes_io.getvalue()

def create_pdf(book: Dict[str, Any]) -> bytes:
    """Create a properly formatted PDF with Times font and content preservation"""
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    
    # Title Page
    pdf.add_page()
    pdf.set_font("Times", "B", 24)
    title = book.get("title", "Untitled Book")
    # Remove any non-ASCII characters
    safe_title = ''.join(c if ord(c) < 128 else '' for c in title).strip()[:150]
    if not safe_title:
        safe_title = "Untitled Book"
    pdf.multi_cell(0, 15, safe_title, align="C")
    
    pdf.ln(10)
    pdf.set_font("Times", "I", 12)
    pdf.cell(0, 8, "Generated by BookForge AI", ln=True, align="C")
    
    pdf.set_font("Times", "", 10)
    timestamp = f"Created on {datetime.now().strftime('%B %d, %Y')}"
    pdf.cell(0, 8, timestamp, ln=True, align="C")
    
    # Table of Contents
    pdf.add_page()
    pdf.set_font("Times", "B", 14)
    pdf.cell(0, 10, "TABLE OF CONTENTS", ln=True)
    pdf.ln(3)
    
    pdf.set_font("Times", "", 11)
    for idx, chapter in enumerate(book.get("chapters", [])[:20], 1):
        chapter_title = chapter.get("title", f"Chapter {idx}")[:70]
        # Remove non-ASCII characters
        safe_chapter = ''.join(c if ord(c) < 128 else '' for c in chapter_title).strip()
        if safe_chapter:
            pdf.cell(0, 6, f"{idx}. {safe_chapter}", ln=True)
    
    # Chapters with content
    for idx, chapter in enumerate(book.get("chapters", [])[:20], 1):
        pdf.add_page()
        
        # Chapter heading
        pdf.set_font("Times", "B", 14)
        chapter_title = chapter.get("title", f"Chapter {idx}")
        safe_title = ''.join(c if ord(c) < 128 else '' for c in chapter_title).strip()
        if safe_title:
            pdf.cell(0, 8, safe_title, ln=True)
        
        # Chapter number
        pdf.set_font("Times", "I", 9)
        pdf.cell(0, 4, f"Chapter {idx}", ln=True)
        pdf.ln(2)
        
        # Content with formatting
        pdf.set_font("Times", "", 11)
        content = chapter.get("content", "")
        if content:
            # Clean content - remove non-ASCII but preserve structure
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(2)
                    continue
                
                # Remove non-ASCII characters but replace bullets with dash
                clean_line = ''
                for c in line:
                    if c == '•' or c == '●' or c == '◦':
                        clean_line += '-'
                    elif ord(c) < 128:
                        clean_line += c
                    elif c in ' \t':
                        clean_line += ' '
                
                clean_line = clean_line.strip()
                if not clean_line:
                    continue
                
                # Detect headers (lines with colon or bold indicators)
                if clean_line.endswith(':') or (len(clean_line) < 80 and clean_line.isupper()):
                    pdf.set_font("Times", "B", 11)
                    pdf.multi_cell(0, 5, clean_line)
                    pdf.ln(1)
                    pdf.set_font("Times", "", 11)
                # Detect bullet points (lines starting with dash)
                elif clean_line.startswith('-') or clean_line.startswith('*'):
                    bullet_text = clean_line.lstrip('-* ').strip()
                    if bullet_text:
                        pdf.multi_cell(0, 5, "- " + bullet_text, l_margin=15)
                        pdf.ln(1)
                # Regular paragraph
                else:
                    pdf.multi_cell(0, 5, clean_line)
                    pdf.ln(1)
            
            pdf.ln(2)
    
    # Output as bytes
    pdf_output = pdf.output()
    if isinstance(pdf_output, str):
        return pdf_output.encode('latin-1', errors='ignore')
    return pdf_output

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "service": "BookForge AI Backend",
        "timestamp": datetime.now().isoformat()
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
